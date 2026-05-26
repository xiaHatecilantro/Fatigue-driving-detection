"""生成优化后的 PPT 文档并保存到原文件"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "微软雅黑"
font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return h


def add_para(text, bold=False, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p


def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "微软雅黑"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p


def add_sep():
    doc.add_paragraph("─" * 60)


# ===== 封面 =====
add_heading("基于计算机视觉的驾驶员疲劳与分心状态检测系统", level=0)
add_para("PPT 文档内容草案（优化版）", bold=True, size=14)
add_para("技术路线：规则法 + 深度学习融合  |  架构：前后端分离  |  部署：笔记本可运行")
doc.add_page_break()

# ===== 01 项目背景 =====
add_heading("01 项目背景", level=1)
add_heading("1.1 行业背景", level=2)
add_para(
    "随着汽车保有量持续增长（全球超 15 亿辆），交通安全问题日益突出。"
    "据 WHO 统计，疲劳驾驶占交通事故成因的 20%-30%，而分心驾驶（低头看手机、调节中控等）"
    "已成为城市交通事故的首要诱因。相比传统人工监管方式，基于计算机视觉的智能监测系统具备"
    "实时性强、部署成本低、可持续运行等优势，是 ADAS（高级驾驶辅助系统）和智能座舱的核心技术之一。"
)
add_heading("1.2 市场与应用趋势", level=2)
add_para(
    "当前智能驾驶辅助、车载安全监控、物流运输监管等领域对驾驶员状态监测需求持续增长。"
    "欧盟已强制要求 2026 年起新车配备驾驶员监测系统（DMS），中国亦在推进相关标准制定。"
    "尤其在网约车、长途货运、校车、危险品运输等场景中，驾驶员状态检测具有极高的实际应用价值。"
    "预计 2028 年全球 DMS 市场规模将突破 50 亿美元。"
)
add_heading("1.3 项目立项意义", level=2)
add_para(
    "本项目围绕“驾驶员疲劳与分心状态检测”展开，目标是设计并实现一个可运行、可展示、可扩展的 Web 系统，"
    "覆盖从数据采集、模型训练、推理融合到前后端展示的完整链路，"
    "为大学生 AI 视觉应用创新项目提供完整原型参考。"
)
add_sep()

# ===== 02 行业痛点 =====
add_heading("02 行业痛点", level=1)
add_heading("2.1 传统方法存在的问题", level=2)
add_para("传统驾驶员状态监测方式主要依赖人工观察、事后回放或单一硬件（红外/脑电）告警，存在以下不足：")
add_bullet("实时性不足 — 人工抽查、事后回看无法及时预警，错失关键干预窗口")
add_bullet("主观性强 — 不同人员判断标准不一，漏报率和误报率均较高")
add_bullet("成本较高 — 专用硬件设备单套数千至数万元，难以大规模部署")
add_bullet("难以量化 — 缺乏统一的疲劳/分心程度评分体系，无法做趋势分析")

add_heading("2.2 视觉检测的实际挑战", level=2)
add_para("虽然视觉方法具备较大潜力，但在真实车载场景中仍面临多项挑战：")
add_bullet("光照变化 — 白天强光、夜间暗光、隧道明暗交替影响人脸检测稳定性")
add_bullet("姿态多样性 — 不同人脸特征、肤色、佩戴眼镜、口罩遮挡等导致关键点提取不稳定")
add_bullet("泛化能力 — 模型在训练集上表现优异，但真实场景中受数据偏置影响，精度下降明显")
add_bullet("实时性要求 — 需要在普通硬件上达到 ≥15 FPS 的检测速率才能实用")

add_heading("2.3 本项目需解决的核心问题", level=2)
add_bullet("如何在普通笔记本/车载边缘设备上实现实时可运行的 Demo")
add_bullet("如何兼顾规则法的可解释性与深度学习模型的增强能力")
add_bullet("如何在双流推理下统一输出风险等级与报警结果")
add_bullet("如何构建前后端分离的完整系统闭环，支持图片/视频/实时三种模式")
add_sep()

# ===== 03 项目概述 =====
add_heading("03 项目概述", level=1)
add_heading("3.1 项目目标", level=2)
add_para("构建一个面向驾驶安全场景的疲劳与分心检测系统，支持三种推理模式：")
add_bullet("本地摄像头实时检测（WebSocket 流式传输，500ms/帧）")
add_bullet("单张图片检测（REST API，multipart/form-data 上传）")
add_bullet("本地视频逐帧分析（REST API，支持逐帧结果与事件分段）")
add_bullet("Web 前端可视化展示（React + TypeScript + Vite）")
add_bullet("统一输出协议：fatigue_score / distraction_score / risk_level / alerts")

add_heading("3.2 检测内容", level=2)
add_para("系统基于 MediaPipe Face Mesh 提取 478 个人脸关键点，重点检测以下驾驶员状态：")
add_bullet("闭眼 — 基于 EAR（Eye Aspect Ratio）计算，阈值 0.22，累积 3 帧触发")
add_bullet("打哈欠 — 基于 MAR（Mouth Aspect Ratio）计算，阈值 0.60，累积 8 帧触发")
add_bullet("头部偏转 — 基于 Pose 关键点近似解算 Yaw/Pitch/Roll，Yaw>18° 或 Pitch>10° 触发")
add_bullet("长时间低头 — Pitch 持续超过阈值 8 帧以上触发")
add_bullet("注意力偏移 — Yaw 大幅偏转 + 模型 distracted 类别协同判断")
add_bullet("可扩展分心行为识别（打电话、喝水、操作中控等）")

add_heading("3.3 系统统一输出协议", level=2)
add_para("所有推理模式共用同一套 UnifiedInferenceResult 数据结构：")
add_bullet("fatigue_score: Float [0, 100] — 疲劳综合评分")
add_bullet("distraction_score: Float [0, 100] — 分心综合评分")
add_bullet("risk_level: \"normal\" | \"mild\" | \"moderate\" | \"severe\" — 风险等级")
add_bullet("rule_result — 规则法独立结果（含风险分、状态标签、原因列表）")
add_bullet("model_result — 模型法独立结果（含预测标签、置信度、四类概率）")
add_bullet("fusion_result — 融合结果（含说明注释、告警列表）")
add_bullet("signals — 底层信号详情（face_detected, ear, mar, yaw, pitch, roll 等）")
add_bullet("alerts — 触发告警标签列表（如 [\"fatigue_warning\", \"distraction_warning\"]）")
add_bullet("timestamp — 时间戳")

add_heading("3.4 系统特点", level=2)
add_bullet("前后端分离 — FastAPI 后端 + React 前端，接口清晰，可独立部署")
add_bullet("双流融合架构 — 规则法主导（70%）+ 模型法增强（30%），兼顾可解释性与准确性")
add_bullet("配置化驱动 — 40+ 参数集中在 YAML 配置文件中，阈值/权重/等级均可调")
add_bullet("轻量模型 — MobileNetV3-Small（~2.5M 参数），适合边缘部署")
add_bullet("时序连续性 — 跨帧窗口累积，非单帧决策，有效降低误报")
add_bullet("事件分段 — 视频推理输出完整事件段（起止时间/持续时长/峰值风险）")
add_sep()

# ===== 04 创新价值 =====
add_heading("04 创新价值", level=1)

add_heading("4.1 算法创新点", level=2)
add_bullet(
    "双流融合引擎（FusionEngine）：规则法 + MobileNetV3 分类模型并行推理，"
    "按 70%:30% 权重加权融合，既保留规则法的可解释性，又引入模型法的非线性纠错能力"
)
add_bullet(
    "模型条件增强策略：当模型高置信度支持规则判断时自动提升分数（model_supported），"
    "当模型高置信度判定正常时进行软抑制（soft_suppression），避免误报"
)
add_bullet(
    "时序状态管理（TemporalState）：跨帧累积闭眼/哈欠/转头/低头的连续帧数，"
    "仅在超过配置阈值后才触发状态变化，大幅降低单帧噪声影响"
)
add_bullet(
    "轻量模型选型：MobileNetV3-Small（~2.5M 参数）在 224x224 输入下实现四分类，"
    "训练时使用 AdamW + ReduceLROnPlateau + EarlyStopping，F1 Macro 达 97.92%"
)
add_bullet(
    "头部姿态近似解算：基于 MediaPipe 5 个关键点（鼻尖/下巴/额头/左脸/右脸），"
    "轻量解算 Yaw/Pitch/Roll，避免引入重量级 PnP 求解器"
)

add_heading("4.2 工程创新点", level=2)
add_bullet(
    "分层架构设计（5 层）：表示层 -> API 网关层 -> 服务编排层 -> 推理逻辑层 -> 信号提取层，"
    "每层职责单一，可独立测试和替换"
)
add_bullet(
    "统一推理协议（UnifiedInferenceResult）：图片/视频/实时三种模式输出相同 JSON Schema，"
    "前端只需一套 ResultCard 组件即可渲染所有结果"
)
add_bullet(
    "配置中心（YAML-driven）：所有阈值、权重、关键点索引、风险等级边界集中在 configs/mvp.yaml，"
    "修改参数无需改动代码，支持快速实验迭代"
)
add_bullet(
    "WebSocket 实时管线：浏览器端 canvas 捕获 -> base64 编码 -> WS 推送 -> 后端解码推理 -> 返回结果，"
    "延迟 < 300ms/帧，实现类实时的 Web 端检测体验"
)
add_bullet(
    "模块化推理服务：InferenceService 统一管理图片/视频/实时三种模式的生命周期，"
    "CommonInferencePipeline 复用核心帧处理逻辑"
)
add_bullet("完善的测试体系：FastAPI TestClient 端到端测试 + pytest 融合引擎单元测试")

add_heading("4.3 应用价值", level=2)
add_bullet("可作为驾驶安全辅助系统原型直接演示，覆盖端到端完整链路")
add_bullet("可扩展到车载 DMS、物流车队监管、驾校培训行为分析、智能座舱交互等场景")
add_bullet("融合架构设计为后续引入更多传感器（红外、方向盘扭矩、生理信号）奠定基础")
add_bullet("轻量模型选型为边缘端部署（Jetson Nano、树莓派等）预留可行性")

add_heading("4.4 比赛展示价值", level=2)
add_bullet("技术路线完整 — 从数据到训练到推理到融合到展示形成闭环")
add_bullet("系统链路清晰 — 5 层架构可分层讲解，每个模块独立可演示")
add_bullet("结果可视化强 — 实时摄像头检测可在比赛现场交互演示")
add_bullet("双流融合是差异化亮点 — 非单一模型堆叠，体现系统性工程思维")
add_sep()

# ===== 05 技术方案（重点丰富） =====
add_heading("05 技术方案", level=1)

add_heading("5.1 总体架构（5 层分层架构）", level=2)

add_para("第一层 — 前端表示层（React 19 + TypeScript 5 + Vite 7）", bold=True)
add_bullet("组件树：App -> TrainingMetricsPanel / UploadPanel / ResultCard / RealtimePanel")
add_bullet("状态管理：React Hooks（useState + useRef）管理推理状态和 WebSocket 连接")
add_bullet("API 通信：api.ts 封装 HTTP REST + WebSocket，统一管理 base URL 和端点")
add_bullet("自定义 Hook：useRealtimeInference 管理 WS 生命周期（connect/disconnect/sendFrame）")
add_bullet("构建工具：Vite 7 开发服务器（HMR 热更新），TypeScript 严格模式编译")

add_para("第二层 — API 网关层（FastAPI 0.136 + Pydantic + Uvicorn 0.46）", bold=True)
add_bullet("REST 端点：GET /health | POST /api/infer/image | POST /api/infer/video | GET /api/metrics/training")
add_bullet("WebSocket 端点：/ws/realtime（JSON 双向通信，帧级实时推理）")
add_bullet("CORS 中间件：允许 localhost:5173 跨域访问")
add_bullet("Pydantic Schema：所有请求/响应用 BaseModel 定义，自动生成 OpenAPI 文档（/docs）")
add_bullet("路由层职责单一：仅做参数解析和响应格式化，零业务逻辑")

add_para("第三层 — 服务编排层（InferenceService / MetricsService）", bold=True)
add_bullet("InferenceService：管理三种推理模式的生命周期（图片/视频/实时）")
add_bullet("TemporaryDirectory：上传文件自动清理，安全的临时文件管理")
add_bullet("帧循环管理：视频推理中的逐帧读取 + 推理 + 事件跟踪")
add_bullet("可选参数开关：save_visualization（保存标注图/视频）、include_frames（返回逐帧结果）")
add_bullet("MetricsService：读取训练产物（best.pt + metrics JSONL）并转换为 API 响应")

add_para("第四层 — 推理逻辑层（CommonInferencePipeline + FusionEngine + ModelRunner）", bold=True)
add_bullet("CommonInferencePipeline：统一单帧处理管线，三种推理模式共享核心逻辑")
add_bullet("处理流程：输入 BGR 帧 -> MediaPipe Face Mesh -> EAR/MAR/HeadPose -> 时序状态 -> 规则评分")
add_bullet("可选模型增强：提取人脸 ROI -> MobileNetV3 推理 -> Softmax -> 四类概率")
add_bullet("FusionEngine：规则分 + 模型概率按 70%:30% 融合，含条件调节（支持/抑制/覆盖）")
add_bullet("ModelRunner：PyTorch 模型推理适配器，封装预处理/推理/后处理")

add_para("第五层 — 信号提取层（cv 模块：EyeFeatures / MouthFeatures / HeadPose / RiskRules）", bold=True)
add_bullet("EyeFeatures：基于 6 个眼睑关键点计算 Eye Aspect Ratio（EAR），阈值 0.22")
add_bullet("MouthFeatures：基于 4 个嘴唇关键点计算 Mouth Aspect Ratio（MAR），阈值 0.60")
add_bullet("HeadPose：基于 5 个 Pose 关键点近似解算 Yaw/Pitch/Roll，含可配置缩放系数")
add_bullet("RiskRules：RuleBasedRiskScorer，接收时序状态后按配置权重累计算分，输出风险等级")

add_heading("5.2 数据流总览", level=2)
add_para(
    "【图片推理】 上传文件 -> 保存临时文件 -> cv2.imread -> "
    "CommonInferencePipeline.process_frame() -> [可选] ModelRunner -> "
    "FusionEngine -> UnifiedInferenceResult -> JSON 响应"
)
add_para(
    "【视频推理】 上传文件 -> VideoCapture 逐帧读取 -> 循环 process_frame() -> "
    "事件跟踪（_update_events）-> 统计聚合 -> VideoInferenceResponse（含摘要 + 事件列表）"
)
add_para(
    "【实时推理】 浏览器 getUserMedia -> canvas.toDataURL() -> WebSocket.send(base64) -> "
    "后端 base64 解码 -> process_frame() -> FusionEngine -> WS.send(JSON) -> 前端渲染"
)

add_heading("5.3 双流融合架构（核心创新）", level=2)
add_para("系统同时运行规则法和模型法两条推理线路，通过 FusionEngine 统一融合：")

add_para("规则法流（70% 权重）：", bold=True)
add_bullet("MediaPipe Face Mesh 提取 478 个 3D 人脸关键点（静态图像模式/视频跟踪模式）")
add_bullet("左眼 6 点 (33,160,158,133,153,144) -> EAR = (|p2-p6|+|p3-p5|) / (2*|p1-p4|)")
add_bullet("右眼 6 点 (362,385,387,263,373,380) -> 同上公式")
add_bullet("嘴部 4 点 (61,13,14,291) -> MAR = |p2-p8| / |p4-p6|")
add_bullet("头姿 5 点 -> 近似映射 Yaw/Pitch/Roll（无需 solvePnP）")
add_bullet("时序累积 -> RuleBasedRiskScorer 按权重累加 -> fatigue_score + distraction_score")

add_para("模型法流（30% 权重）：", bold=True)
add_bullet("模型：MobileNetV3-Small（torchvision），分类头 576->1024->4，约 2.5M 参数")
add_bullet("输入：224x224 RGB 人脸 ROI（或全帧），ImageNet 标准化")
add_bullet("输出：Softmax -> {normal, eye_closed, yawn, distracted} 四类概率")
add_bullet("分数映射：eye_closed_prob x100, yawn_prob x85, distracted_prob x100")

add_para("融合策略：", bold=True)
add_bullet("综合分 = rule_score x0.7 + model_mapped_score x0.3")
add_bullet("模型高置信度支持规则时 -> 自动提升分数 + 标记 model_supported")
add_bullet("模型高置信度判正常（normal>=0.90）且规则无强异常 -> 软抑制 x0.9")
add_bullet("最终映射到 risk_level：mild(30-60), moderate(60-80), severe(80-100)")

add_heading("5.4 配置化系统设计", level=2)
add_para("所有运行参数集中在 configs/mvp.yaml，零硬编码：")
add_bullet("camera — 摄像头索引与分辨率 (640x480)")
add_bullet("mediapipe — Face Mesh 置信度、最大人脸数、精细化标志")
add_bullet("landmarks — 左眼/右眼 6 点、嘴部 4 点、头姿 5 点的 MediaPipe 索引")
add_bullet("thresholds — EAR 闭眼阈值 (0.22)、MAR 哈欠阈值 (0.60)")
add_bullet("head_pose — Yaw(18度)/Pitch(10度) 阈值及缩放系数")
add_bullet("temporal — 闭眼(3帧)/哈欠(8帧)/转头(8帧)/低头(8帧) 触发帧数")
add_bullet("risk — 告警阈值(60)、等级边界、各信号的基础分/步进分/最大提升分")
add_bullet("fusion — 规则权重(0.7)、模型权重(0.3)、支持阈值、抑制阈值、分数映射")
add_bullet("训练配置：training/configs/base.yaml（学习率/批次/早停/增强）")
add_bullet("数据集配置：training/configs/dataset_map.yaml（数据路径与类别映射）")

add_heading("5.5 完整技术栈", level=2)

add_para("后端技术栈：", bold=True)
add_bullet("语言：Python 3.12")
add_bullet("Web 框架：FastAPI 0.136（异步 ASGI）+ Pydantic（数据校验）")
add_bullet("ASGI 服务器：Uvicorn 0.46（支持热重载开发）")
add_bullet("计算机视觉：OpenCV (cv2) -- 图像读取/视频解码/帧处理")
add_bullet("人脸关键点：MediaPipe Face Mesh -- 478 点 3D 网格")
add_bullet("数值计算：NumPy -- 数组运算与矩阵变换")
add_bullet("深度学习：PyTorch + torchvision（MobileNetV3-Small 预训练权重）")
add_bullet("配置管理：PyYAML -- YAML 文件解析")
add_bullet("测试：FastAPI TestClient + pytest")

add_para("前端技术栈：", bold=True)
add_bullet("框架：React 19.1（函数组件 + Hooks）")
add_bullet("类型系统：TypeScript 5.8（严格模式）")
add_bullet("构建工具：Vite 7.0（HMR 开发服务器 + 生产构建）")
add_bullet("编译插件：@vitejs/plugin-react 5.0")
add_bullet("样式：原生 CSS（global.css），零额外依赖")

add_para("训练技术栈：", bold=True)
add_bullet("模型：MobileNetV3-Small（约 2.5M 参数，ImageNet 预训练）")
add_bullet("分类头：Linear(576->1024) + Hardswish + Dropout(0.3) + Linear(1024->4)")
add_bullet("优化器：AdamW（lr=5e-4, weight_decay=5e-4）")
add_bullet("调度器：ReduceLROnPlateau（factor=0.5, patience=2）")
add_bullet("早停：patience=4 epochs，按 F1 Macro 保存最佳 checkpoint")
add_bullet("数据增强：RandomHorizontalFlip(0.5) + ColorJitter + RandomRotation(±8度)")
add_bullet("训练日志：JSONL（每行一个 epoch 的完整指标）")

add_para("通信协议：", bold=True)
add_bullet("HTTP REST：multipart/form-data（文件上传）+ JSON（响应）")
add_bullet("WebSocket：JSON 双向帧级通信（实时推理）")
add_bullet("CORS：允许 http://localhost:5173 跨域")

add_heading("5.6 系统文件架构", level=2)
add_para("项目源码按功能模块组织，共约 40 个核心源文件：")
add_bullet("backend/（12 文件）-- FastAPI 入口、路由、Schema、Service 编排、推理服务管理")
add_bullet("frontend/src/（10 文件）-- App、4 个组件、API 客户端、自定义 Hook、类型定义")
add_bullet("cv/（6 文件）-- 眼/嘴/头姿特征提取、规则评分引擎")
add_bullet("inference/（7 文件）-- 统一管线、融合引擎、模型运行器、三种推理 CLI")
add_bullet("training/（6 文件）-- 训练/评估入口、MobileNetV3 模型、自定义数据集、训练配置")
add_bullet("configs/（1 文件）-- mvp.yaml（40+ 运行时参数）")
add_bullet("scripts/（3 文件）-- 数据集准备、ROI 裁剪、分组划分")
add_bullet("tests/（2 文件）-- API 端到端测试、融合引擎单元测试")
add_sep()

# ===== 06 成果验证 =====
add_heading("06 成果验证", level=1)

add_heading("6.1 已完成的系统功能", level=2)
add_bullet("前后端联调 -- React 前端 <-> FastAPI 后端完整通路")
add_bullet("图片推理接口 -- POST /api/infer/image，支持可选可视化保存")
add_bullet("视频推理接口 -- POST /api/infer/video，返回逐帧结果 + 事件分段 + 统计摘要")
add_bullet("实时 WebSocket 推理 -- /ws/realtime，500ms/帧，< 300ms 延迟")
add_bullet("训练成果展示页 -- TrainingMetricsPanel 展示混淆矩阵、训练曲线、类别指标")
add_bullet("规则法 / 模型法 / 融合结果 -- 三层结果同屏展示，可解释性透明")
add_bullet("健康检查 + API 文档 -- /health 端点 + 自动生成的 /docs 交互式文档")

add_heading("6.2 模型训练结果", level=2)
add_para("基于 MobileNetV3-Small，使用 ROI 裁剪版数据集（4 类样本）训练，baseline 评估：")
add_bullet("Accuracy（准确率）：97.80%")
add_bullet("Precision Macro（宏精确率）：98.25%")
add_bullet("Recall Macro（宏召回率）：97.63%")
add_bullet("F1 Macro（宏 F1）：97.92%")
add_para(
    "训练配置：AdamW (lr=5e-4, wd=5e-4) + ReduceLROnPlateau + EarlyStopping(patience=4) + "
    "数据增强（翻转/色彩抖动/旋转），按 F1 Macro 保存 best.pt。"
)

add_heading("6.3 推理性能", level=2)
add_bullet("单帧推理耗时（不含模型）：约15-25ms（MediaPipe + EAR/MAR/HeadPose + 规则评分）")
add_bullet("单帧推理耗时（含模型）：约30-50ms（额外增加 MobileNetV3 前向传播）")
add_bullet("实时帧率：>=20 FPS（笔记本 CPU），>=30 FPS（GPU 加速）")

add_heading("6.4 融合效果分析", level=2)
add_bullet("eye_closed、yawn 类别模型与规则法一致性高，互相增强效果好")
add_bullet("distracted 类别模型提供额外检测能力，弥补规则法对复杂分心行为覆盖不足")
add_bullet("规则法在实时性上表现稳定（无模型推理时延），适合持续监控场景")
add_bullet("模型法增强后，系统整体判断比单独规则法更合理，尤其在边界情况下")
add_bullet("融合引擎的条件调节机制（软抑制/模型支持）有效降低误报率")

add_heading("6.5 当前局限与改进方向", level=2)
add_bullet("数据来源分布不均 -> distracted 类别样本偏少，泛化能力需增强")
add_bullet("图片与实时场景域差异 -> 模型概率在真实光照/姿态变化下可能波动")
add_bullet("关键点检测 -> 复杂光照和大角度遮挡（墨镜、口罩）下仍不稳定")
add_bullet("硬件依赖 -> 实时推理对 CPU/GPU 有一定要求（>=4 核 CPU 推荐）")
add_sep()

# ===== 07 团队分工 =====
add_heading("07 团队分工", level=1)

add_heading("7.1 算法方向", level=2)
add_bullet("规则法设计与实现（EAR/MAR/HeadPose 特征提取 + RiskRules 评分引擎）")
add_bullet("模型训练与评估（MobileNetV3-Small 四分类 + 数据增强 + 超参调优）")
add_bullet("融合策略设计（FusionEngine 权重分配 + 条件调节逻辑）")
add_bullet("时序状态管理（TemporalState 帧累积 + 事件分段算法）")

add_heading("7.2 后端方向", level=2)
add_bullet("FastAPI 接口设计与实现（REST 路由 + WebSocket 端点 + Pydantic Schema）")
add_bullet("图片/视频/实时三种推理模式服务封装（InferenceService）")
add_bullet("统一输出协议设计（UnifiedInferenceResult）与 CommonInferencePipeline 整合")
add_bullet("临时文件管理、配置加载、健康检查和 API 文档")

add_heading("7.3 前端方向", level=2)
add_bullet("可视化页面搭建（UploadPanel / ResultCard / RealtimePanel / TrainingMetricsPanel）")
add_bullet("图片/视频上传与结果展示（含规则/模型/融合三层结果同屏渲染）")
add_bullet("实时检测界面（WebSocket 连接 + canvas 帧捕获 + 实时结果更新）")
add_bullet("训练成果展示（混淆矩阵 + 训练曲线 + 类别指标）")

add_heading("7.4 项目统筹", level=2)
add_bullet("整体架构设计（5 层分层架构 + 双流融合方案）")
add_bullet("配置化体系搭建（YAML 配置 + 模块化参数管理）")
add_bullet("里程碑管理与版本推进")
add_bullet("答辩材料组织与技术亮点提炼")
add_sep()

# ===== 08 未来展望 =====
add_heading("08 未来展望", level=1)

add_heading("8.1 算法优化方向", level=2)
add_bullet("增加真实车载场景数据采集与标注（不同光照/姿态/遮挡条件）")
add_bullet("探索多帧时序模型（如轻量 Transformer/LSTM）替代单帧分类，提升时序一致性")
add_bullet("优化关键点检测稳定性（引入红外图像融合、光照增强预处理）")
add_bullet("扩展细粒度分心行为识别（打电话、喝水、操作中控、转身交谈等）")
add_bullet("引入对比学习和数据增强策略提升模型泛化能力")

add_heading("8.2 系统优化方向", level=2)
add_bullet("前端静态部署（Vite build -> Nginx/CDN）+ 后端常驻运行（systemd/Docker）")
add_bullet("增加趋势图（历史疲劳/分心分数折线图）、事件时间轴、历史统计面板")
add_bullet("增加一键启动脚本（start.bat / start.sh），适配比赛现场快速部署")
add_bullet("增加日志系统（structlog）、Prometheus 监控指标、异常恢复机制")
add_bullet("模型热替换：支持运行时切换不同 checkpoint 做 A/B 对比")
add_bullet("Docker 容器化部署，统一开发和生产环境")

add_heading("8.3 应用扩展方向", level=2)
add_bullet("车载终端适配 -- 接入车载摄像头实时流（RTSP），适配嵌入式平台（Jetson Nano/树莓派）")
add_bullet("多传感器融合 -- 融合方向盘扭矩传感器、红外眼动仪等硬件信号")
add_bullet("车队管理平台 -- 多车实时监控大盘 + 驾驶员画像 + 安全评分")
add_bullet("驾驶培训 -- 学员行为分析 + 教练实时指导 + 培训报告自动生成")
add_bullet("智能座舱交互 -- 基于驾驶员状态的 HMI 自适应（疲劳时自动提醒休息）")

add_heading("8.4 总结", level=2)
add_para(
    "本项目已经完成从算法设计、模型训练、双流融合推理到前后端展示的完整原型实现，"
    "技术路线清晰、架构设计规范、代码模块化程度较高。"
    "项目在双流融合架构上的设计体现了系统化工程思维，而非简单的模型堆叠。"
    "后续将继续围绕数据质量、模型泛化和工程稳定性进行优化，"
    "逐步提升系统的实用性和应用落地能力。"
)

# ── 保存 ──
output_path = r"C:\Users\Xia chuan can\Desktop\新建 DOCX 文档.docx"
doc.save(output_path)
print(f"已保存到: {output_path}")
